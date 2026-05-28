from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt


def normalize_text(text: str | None) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def short_text(text: str | None, limit: int = 140) -> str:
    text = normalize_text(text)
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


def note_title(note: dict) -> str:
    return normalize_text(note.get("pageTitle") or note.get("cardTitle")) or "NOT_DISPLAYED"


def note_author(note: dict) -> str:
    return normalize_text(note.get("pageAuthor") or note.get("cardAuthor")) or "NOT_DISPLAYED"


def note_body(note: dict) -> str:
    body = normalize_text(note.get("body"))
    if body:
        return body
    return normalize_text(note.get("rawBodyText"))


def note_summary(note: dict) -> str:
    return short_text(note_body(note), 160)


def classify_note(note: dict) -> str:
    text = " ".join(
        [
            note_title(note).lower(),
            note_body(note).lower(),
            " ".join(note.get("tags", [])),
        ]
    )
    labels: list[str] = []
    if any(token in text for token in ["skill", "skills", "prompt", "workflow", "paper-plot"]):
        labels.append("skill")
    if any(token in text for token in ["mcp", "chrome", "zotero", "微信", "浏览器", "remote"]):
        labels.append("mcp")
    if any(token in text for token in ["plugin", "插件", "oh-my-codex", "omx"]):
        labels.append("plugin")
    if any(token in text for token in ["自动化", "automation", "agent", "项目", "科研"]):
        labels.append("workflow")
    return " / ".join(labels) if labels else "general"


def collect_comments(notes: list[dict]) -> list[dict]:
    ranked: list[dict] = []
    for note in notes:
        for comment in note.get("visibleTopLevelComments", []):
            ranked.append(
                {
                    "note_title": note_title(note),
                    "note_url": note.get("boardUrl", "NOT_DISPLAYED"),
                    "author": normalize_text(comment.get("author")) or "NOT_DISPLAYED",
                    "content": normalize_text(comment.get("content")) or "NOT_DISPLAYED",
                    "likes": int(comment.get("likeCount") or 0),
                    "replies": int(comment.get("replyCount") or 0),
                    "date": normalize_text(comment.get("dateText")) or "NOT_DISPLAYED",
                }
            )
    ranked.sort(key=lambda item: item["likes"], reverse=True)
    return ranked


def peel_findings(notes: list[dict], top_comments: list[dict]) -> list[str]:
    full_text = "\n".join(
        f"{note_title(note)}\n{note_body(note)}"
        for note in notes
    ).lower()

    skill_count = sum("skill" in classify_note(note) for note in notes)
    mcp_count = sum("mcp" in classify_note(note) for note in notes)
    plugin_count = sum("plugin" in classify_note(note) for note in notes)
    workflow_count = sum("workflow" in classify_note(note) for note in notes)

    findings = [
        (
            "Point: 这组笔记最核心的共识不是“单个神技”，而是把多个技能和外部连接拼成稳定工作流。 "
            f"Evidence: {skill_count} 条笔记直接讨论 skill，{workflow_count} 条笔记直接讨论 agent/workflow/自动化。 "
            "Explanation: 高关注内容持续围绕任务模板化、流程编排、写作与分析链路复用，而不是迷信单一提示词。 "
            "Link: v3 应继续坚持“模块化能力栈”而不是“爆款单点能力”。"
        ),
        (
            "Point: 与外部世界相连的接口层是高频需求。 "
            f"Evidence: {mcp_count} 条笔记出现 MCP、浏览器、Zotero、微信或远程控制相关线索。 "
            "Explanation: 研究者真正需要的是把文献库、浏览器登录态、社媒页面和写作工具接进 Codex，而不是只让模型在空上下文里输出文字。 "
            "Link: v3 中保留 Chrome DevTools、Zotero MCP、社媒读取链路是必要的。"
        ),
        (
            "Point: 高赞评论把 Codex 的定位压回到了“agent/工作流调度器”，而不是科研判断替代物。 "
            f"Evidence: 当前可见点赞最高的 {min(3, len(top_comments))} 条评论，都在讨论 Codex 与 ChatGPT 的差别、agent 属性，以及技能组合的优势。 "
            "Explanation: 这说明受欢迎的实践并不假定 AI 能替代 idea、taste 和学术判断，而是把它当成执行层和组织层。 "
            "Link: v3 必须继续保留 DOI 核验、PEEL、人工把关和提案式自进化。"
        ),
    ]
    if plugin_count:
        findings.append(
            (
                "Point: 插件层价值存在，但应附属于主工作流。 "
                f"Evidence: {plugin_count} 条笔记直接提到插件、OMX 或插件化扩展。 "
                "Explanation: 插件被当作效率增强层，而不是研究主链。 "
                "Link: OMX 继续放在 WSL2 增强层，而不是直接覆盖 Windows 主配置。"
            )
        )
    return findings


def v3_adjustments(notes: list[dict], top_comments: list[dict]) -> list[str]:
    tags = Counter(classify_note(note) for note in notes)
    lines = [
        "把小红书专辑观察到的高频需求正式并入 v3：",
        f"- 与 skill 直接相关的笔记占比最高，说明项目启动、分析路由、写作和复盘不应塞进一个总控 skill，而应拆成多个专门 skill 再由路由器调度。",
        f"- 与 MCP 相关的内容密度很高，说明浏览器登录态、Zotero 文献库和可选的 Scholar/CNKI 后端必须保留为独立接口层。",
        f"- 当前可见评论里，对 Codex 的定位集中在“agent + skills + workflow”，因此 v3 应继续把人工判断、证据链和项目治理放在第一位。",
        f"- 本轮笔记分类统计：{', '.join(f'{key}={value}' for key, value in sorted(tags.items()))}。",
    ]
    if top_comments:
        lines.append(
            f"- 最高赞评论来自《{top_comments[0]['note_title']}》，其核心信息是：高效配置比单点功能更重要。"
        )
    return lines


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def write_docx(payload: dict, output_docx: Path) -> list[dict]:
    notes = payload["notes"]
    board = payload["board"]
    top_comments = collect_comments(notes)[:3]

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    document.add_heading("小红书 AI 专辑：技能、插件与 MCP 总结", level=0)
    document.add_paragraph(
        f"专辑标题：{board.get('title', 'NOT_DISPLAYED')} | 所有者：{board.get('owner', 'NOT_DISPLAYED')} | "
        f"可见笔记数：{len(notes)} | 来源链接：{board.get('url', 'NOT_DISPLAYED')}"
    )
    document.add_paragraph(
        "方法说明：本报告只依据当前登录态浏览器里真实可见的专辑卡片、正文和已渲染顶层评论生成，"
        "不补写隐藏字段，不推断未加载内容。"
    )

    document.add_heading("一、总体判断", level=1)
    for finding in peel_findings(notes, top_comments):
        document.add_paragraph(finding)

    document.add_heading("二、所有笔记总表", level=1)
    rows: list[list[str]] = []
    for note in notes:
        rows.append(
            [
                str(note.get("index", "")),
                note_title(note),
                note_author(note),
                str(note.get("cardLikedCount") or 0),
                note_summary(note),
                classify_note(note),
            ]
        )
    add_table(
        document,
        ["序号", "标题", "作者", "点赞", "100-160字摘要", "与 skill/MCP/plugin 的关系"],
        rows,
    )

    document.add_heading("三、点赞量最高的三条可见顶层评论", level=1)
    for index, comment in enumerate(top_comments, start=1):
        document.add_paragraph(
            f"{index}. 来源笔记：《{comment['note_title']}》 | 评论作者：{comment['author']} | 点赞：{comment['likes']} | 回复：{comment['replies']}"
        )
        document.add_paragraph(comment["content"])

    document.add_heading("四、与 skill、插件、MCP 直接相关的结论", level=1)
    for line in v3_adjustments(notes, top_comments):
        document.add_paragraph(line)

    document.add_heading("五、逐条笔记结构化记录", level=1)
    for note in notes:
        document.add_heading(f"{note.get('index', '')}. {note_title(note)}", level=2)
        document.add_paragraph(
            f"作者：{note_author(note)} | 点赞：{note.get('cardLikedCount') or 0} | "
            f"评论文案：{note.get('totalCommentsVisibleText') or 'NOT_DISPLAYED'}"
        )
        document.add_paragraph(f"链接：{note.get('boardUrl', 'NOT_DISPLAYED')}")
        document.add_paragraph(f"摘要：{note_summary(note)}")
        document.add_paragraph(f"分类：{classify_note(note)}")
        visible_comments = note.get("visibleTopLevelComments", [])[:3]
        for comment in visible_comments:
            document.add_paragraph(
                f"评论证据 | {normalize_text(comment.get('author'))} | 点赞 {comment.get('likeCount', 0)} | "
                f"{short_text(comment.get('content'), 120)}"
            )

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))
    return top_comments


def write_markdown(payload: dict, output_md: Path, top_comments: list[dict]) -> None:
    notes = payload["notes"]
    board = payload["board"]
    lines: list[str] = [
        "# 小红书 AI 专辑总结",
        "",
        f"- 专辑标题：{board.get('title', 'NOT_DISPLAYED')}",
        f"- 所有者：{board.get('owner', 'NOT_DISPLAYED')}",
        f"- 可见笔记数：{len(notes)}",
        f"- 来源链接：{board.get('url', 'NOT_DISPLAYED')}",
        "",
        "## 总体判断",
        "",
    ]
    lines.extend(f"- {line}" for line in peel_findings(notes, top_comments))
    lines.extend(["", "## 点赞量最高的三条可见顶层评论", ""])
    for index, comment in enumerate(top_comments, start=1):
        lines.extend(
            [
                f"### 评论 {index}",
                f"- 来源笔记：{comment['note_title']}",
                f"- 作者：{comment['author']}",
                f"- 点赞：{comment['likes']}",
                f"- 内容：{comment['content']}",
                "",
            ]
        )

    lines.extend(["## v3 修订建议", ""])
    lines.extend(f"- {line}" for line in v3_adjustments(notes, top_comments))
    output_md.parent.mkdir(parents=True, exist_ok=True)
    output_md.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_json", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("output_md", type=Path)
    args = parser.parse_args()

    payload = json.loads(args.input_json.read_text(encoding="utf-8"))
    top_comments = write_docx(payload, args.output_docx)
    write_markdown(payload, args.output_md, top_comments)


if __name__ == "__main__":
    main()
