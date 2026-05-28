from __future__ import annotations

import html
import json
import math
import subprocess
import shutil
import sys
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from PIL import Image, ImageDraw, ImageFont

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Inches
except Exception:  # pragma: no cover - docx is optional at runtime
    Document = None
    Inches = None


REPO_ROOT = Path(__file__).resolve().parents[2]
SKILLS_ROOT = REPO_ROOT / "skills"
CATALOG_ROOT = SKILLS_ROOT / "catalog"
OUTPUT_ROOT = REPO_ROOT / "outputs" / "environment-overview"
IMAGE_ROOT = OUTPUT_ROOT / "images"
PRESENTATION_ROOT = OUTPUT_ROOT / "presentation"
ASSET_ROOT = REPO_ROOT / "assets"
IMAGE2_OVERRIDE_ROOT = ASSET_ROOT / "environment-overview-image2"
GUIZANG_ROOT = Path.home() / ".codex" / "skills" / "guizang-ppt-skill"

PAPER = "#FAFAF8"
INK = "#0A0A0A"
MUTED = "#5F6673"
LINE = "#DCDCDC"
GREY = "#F0F0EE"
BLUE = "#002FA7"
TEAL = "#0F9F8F"
PURPLE = "#6E56CF"
ORANGE = "#F59E0B"
GREEN = "#17A673"


SKILL_PURPOSES = {
    "research-autopilot": "总入口。先判断任务路线、工具组合和是否需要项目协作。",
    "research-team-orchestrator": "项目型任务分工。把执行、审查、产物和交接关系写清楚。",
    "research-stack-manager": "环境治理。检查 skill、插件、工具接口、路径和配置漂移。",
    "project-retrospective-evolver": "项目复盘。把项目经验转成可审查的环境改进候选。",
    "evidence-based-literature-workflow": "证据核验总控。结构性阅读、候选文献、引文支撑和 PDF 标注优先进入这里。",
    "reference-fulltext-acquisition": "论文全文获取。统一调度 OpenAlex、Google Scholar、CNKI、浏览器下载和 PDF 校验。",
    "citation-verifier": "正式引用核验。检查作者、年份、题名、来源和 DOI 或替代证据。",
    "zotero-sync": "Zotero 同步。把核验过的文献、附件、标签和收藏夹写入正式文献库。",
    "google-scholar-research": "Google Scholar 检索。逐篇找论文、右侧 PDF 和引用线索。",
    "cnki-research": "中文学术检索。做 CNKI 候选发现、详情页证据和下载前核验。",
    "openalex-landscape": "开放学术版图。查主题、作者、机构、期刊和引用网络。",
    "semantic-citation-tracer": "语义引文追踪。找相近论文、上游引用和下游影响。",
    "systematic-literature-review": "系统综述。整理多篇文献的主题、争论、证据和研究空白。",
    "pdf": "PDF 处理。读全文、看版面、抽取正文和核验文件质量。",
    "academic-paper-review": "单篇论文评审。检查贡献、方法、证据、论证和修改方向。",
    "quant-analysis": "量化研究。处理回归、DID、IV、RDD、稳健性和结果解释。",
    "text-analysis": "文本分析。处理语料清洗、编码、主题、词典和嵌入。",
    "network-analysis": "网络分析。处理节点、边、中心性、社群和扩散。",
    "research-design-studio": "研究设计。把想法转成问题、对象、材料和方法边界。",
    "dataset-discovery": "数据发现。寻找开放数据、复现数据和平台导出。",
    "digital-trace-pipeline": "数字痕迹材料。处理平台、网页和社媒证据边界。",
    "abm-simulation-lab": "仿真实验。设计代理人规则、参数和敏感性分析。",
    "reproducibility-package": "复现包。整理脚本、数据、环境说明和可复跑路径。",
    "long-running-experiment-ops": "长时任务运维。监控后台任务、断点续跑和失败恢复。",
    "figure-table-studio": "结果图表。做数据图、回归表和论文级结果展示。",
    "research-figure-studio": "科研绘图。做机制图、概念图、多面板图和投稿级图形审查。",
    "research-presentation-studio": "研究演示。生成 PPT、网页 PPT、答辩和汇报材料。",
    "guizang-ppt-skill": "高质量网页 PPT。提供电子杂志风和瑞士国际主义两套视觉系统。",
    "manuscript-writing-studio": "论文写作。处理正文、改写、润色、目标期刊适配和去模板化表达。",
    "academic-humanization-studio": "学术去模板感。只改表层语言，不改变事实、数据、引用和结论。",
    "writing-reference-capture": "写作引用捕获。只记录文稿中真实使用的文献和证据句。",
    "reviewer-response-pack": "审稿回复包。整理审稿意见、回应矩阵和修改位置。",
    "social-science-submission-packager": "社科投稿包。冻结文稿、图表、附录、引用和复现材料。",
    "latex-paper-conversion": "LaTeX 模板迁移。把论文从一个模板迁到另一个模板。",
    "research-docx-export": "Word 导出。把报告、评审、说明和项目材料整理成 docx。",
    "scholar-nuwa": "学者蒸馏。把用户确认的学者材料做成可调用角色。",
    "scholar-panel": "多学者讨论。让蒸馏学者独立审阅、互相回应并形成共识。",
    "obsidian-research-sync": "Obsidian 沉淀。同步阅读笔记、项目复盘和方法卡。",
    "social-platform-reader": "平台材料读取。保存浏览器可见证据、截图和来源信息。",
    "local-cloud-router": "本地/云端判断。决定任务适合本机还是云端执行。",
    "skill-vetter": "外部能力审查。引入第三方仓库、skill、插件前先评估价值和风险。",
    "playwright": "浏览器自动化。做可复现网页操作、截图、下载和界面检查。",
    "playwright-interactive": "交互式浏览器调试。快速观察网页或桌面界面状态。",
    "agent-browser": "浏览器辅助。配合平台证据读取和页面交互。",
}

MCP_PURPOSES = {
    "chrome-devtools": "连接 Chrome，用于页面检查、截图、网络观察和浏览器证据。",
    "playwright-mcp": "控制真实浏览器，做可复现网页操作和下载流程。",
    "zotero-mcp": "连接 Zotero，管理文献条目、附件、笔记和收藏夹。",
    "openalex-mcp": "连接 OpenAlex，做开放学术检索和引用关系查询。",
    "semantic-scholar-mcp": "连接 Semantic Scholar，找相近论文和引用线索。",
    "google-scholar-mcp": "连接 Google Scholar 检索链，辅助逐篇找论文和 PDF。",
    "cnki-mcp": "中文学术检索观察项，用于候选发现和页面证据。",
    "paper-search-mcp": "聚合论文搜索，用于补充发现。",
    "social-platform-mcp": "平台材料统一采集接口，保存浏览器可见证据。",
    "xiaohongshu-mcp": "小红书专用后端，不是平台总入口。",
    "figma-dev-mode-mcp": "连接 Figma 设计上下文，服务桌面应用开发。",
    "agentmemory": "运行态记忆召回和审计，只能作为候选记忆层。",
    "codegraph": "代码结构索引，只能辅助理解影响范围。",
}


def load_json(name: str) -> dict[str, Any]:
    return json.loads((CATALOG_ROOT / name).read_text(encoding="utf-8-sig"))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size=size)
    return ImageFont.load_default()


def draw_round(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str = LINE, radius: int = 28, width: int = 2) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], value: str, size: int, fill: str = INK, bold: bool = False, anchor: str | None = None) -> None:
    draw.text(xy, value, font=font(size, bold), fill=fill, anchor=anchor)


def wrap(draw: ImageDraw.ImageDraw, value: str, max_width: int, size: int, bold: bool = False) -> list[str]:
    result: list[str] = []
    current = ""
    fnt = font(size, bold)
    for ch in value:
        candidate = current + ch
        if draw.textlength(candidate, font=fnt) <= max_width or not current:
            current = candidate
        else:
            result.append(current)
            current = ch
    if current:
        result.append(current)
    return result


def paragraph(draw: ImageDraw.ImageDraw, xy: tuple[int, int], value: str, max_width: int, size: int, fill: str = MUTED, line_gap: int = 8) -> int:
    y = xy[1]
    for line in wrap(draw, value, max_width, size):
        text(draw, (xy[0], y), line, size, fill=fill)
        y += size + line_gap
    return y


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = BLUE, width: int = 5) -> None:
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    pts = [
        end,
        (int(end[0] - size * math.cos(angle - 0.45)), int(end[1] - size * math.sin(angle - 0.45))),
        (int(end[0] - size * math.cos(angle + 0.45)), int(end[1] - size * math.sin(angle + 0.45))),
    ]
    draw.polygon(pts, fill=fill)


def save(img: Image.Image, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, "PNG", optimize=True)


def draw_main_map(path: Path) -> None:
    img = Image.new("RGB", (2600, 1500), PAPER)
    d = ImageDraw.Draw(img)
    text(d, (90, 70), "本地研究环境", 60, bold=True)
    text(d, (92, 145), "以工程控制论组织研究、工具、证据和长期演化", 30, fill=MUTED)

    lanes = [
        ("任务入口", BLUE),
        ("规则底座", TEAL),
        ("执行工作", PURPLE),
        ("工具与记忆", ORANGE),
    ]
    y0 = 270
    for i, (label, color) in enumerate(lanes):
        y = y0 + i * 265
        d.line([(80, y + 90), (2450, y + 90)], fill="#ECEDEF", width=2)
        d.ellipse((52, y + 66, 96, y + 110), fill=color)
        text(d, (116, y + 70), label, 28, bold=True)

    draw_round(d, (250, 235, 720, 395), "#FFFFFF", "#CFD6E3")
    text(d, (295, 282), "用户任务 / Codex 会话", 34, bold=True)
    paragraph(d, (295, 330), "研究、写作、文献、数据分析、绘图、PPT、环境维护。", 355, 22)

    draw_round(d, (875, 220, 1380, 430), BLUE, BLUE)
    text(d, (930, 275), "总入口：先判断路线", 38, "#FFFFFF", True)
    paragraph(d, (930, 330), "research-autopilot 先看任务、材料、边界和风险；新链条或语义模糊时先询问。", 395, 22, "#EAF1FF")
    arrow(d, (720, 315), (875, 315))

    draw_round(d, (1580, 225, 2185, 405), "#F0FBF8", "#8BD5C8")
    text(d, (1620, 270), "控制核与规则", 36, bold=True)
    paragraph(d, (1620, 322), "目标、边界、反馈信号、稳定性指标、记忆准入和冲突处理。", 460, 23)
    arrow(d, (1380, 315), (1580, 315))

    layer_cards = [
        ("执行动作", "运行命令、写文件、生成结果", BLUE),
        ("工具接口", "浏览器、Zotero、学术检索", ORANGE),
        ("上下文与证据", "材料、引用、项目事实", TEAL),
        ("研究阶段", "从设计到投稿和复盘", PURPLE),
        ("运行日志", "日报、状态、失败原因", "#64748B"),
        ("可靠性检查", "结构规则、测试、检查点", "#7C3AED"),
        ("环境治理", "权限、回滚、外部吸收", GREEN),
    ]
    start_x = 250
    for idx, (title, body, color) in enumerate(layer_cards):
        x = start_x + (idx % 4) * 530
        y = 520 + (idx // 4) * 150
        draw_round(d, (x, y, x + 430, y + 118), "#FFFFFF", "#D7DDE8", 22)
        d.rectangle((x, y, x + 12, y + 118), fill=color)
        text(d, (x + 36, y + 25), title, 28, bold=True)
        paragraph(d, (x + 36, y + 66), body, 330, 20)

    work_cards = [
        ("找文献", "候选池 / PDF / Zotero"),
        ("结构性阅读", "证据句 / 文献综述"),
        ("分析与图表", "量化 / 文本 / 网络 / 绘图"),
        ("写作与审稿", "正文 / 回复 / 润色"),
        ("PPT 与交付", "汇报 / 投稿包 / 复现包"),
    ]
    for idx, (title, body) in enumerate(work_cards):
        x = 250 + idx * 410
        draw_round(d, (x, 845, x + 330, 1015), "#FFFFFF", "#DADFF0", 24)
        text(d, (x + 34, 895), title, 30, bold=True)
        paragraph(d, (x + 34, 945), body, 248, 21)

    bottom_cards = [
        ("Zotero", "正式文献库"),
        ("Obsidian", "长解释与复盘"),
        ("agentmemory", "运行态记忆召回"),
        ("codegraph", "代码结构索引"),
        ("Git", "版本记录与回滚"),
    ]
    for idx, (title, body) in enumerate(bottom_cards):
        x = 250 + idx * 410
        draw_round(d, (x, 1120, x + 330, 1265), "#FFFDF7", "#ECD8A8", 22)
        text(d, (x + 34, 1160), title, 28, bold=True)
        paragraph(d, (x + 34, 1210), body, 255, 20)

    draw_round(d, (2265, 540, 2515, 1020), "#F7F5FF", "#B8ADF2", 26)
    text(d, (2300, 585), "检查反馈", 30, bold=True)
    checks = ["路线是否选对", "证据是否足够", "工具是否越界", "结果能否复查", "经验是否沉淀"]
    for i, item in enumerate(checks):
        y = 650 + i * 66
        d.ellipse((2300, y, 2325, y + 25), fill=PURPLE)
        text(d, (2345, y - 4), item, 23)
    for x in range(2280, 1420, -38):
        d.line([(x, 445), (x - 18, 445)], fill=BLUE, width=4)
    arrow(d, (1420, 445), (1380, 360), fill=BLUE, width=4)
    text(d, (1510, 420), "闭环反馈与纠偏", 22, fill=BLUE, bold=True)
    save(img, path)


def draw_stage_map(path: Path) -> None:
    img = Image.new("RGB", (2400, 1450), PAPER)
    d = ImageDraw.Draw(img)
    text(d, (90, 70), "研究工作从哪里开始、下一步去哪里", 54, bold=True)
    text(d, (92, 140), "AI 会主动判断当前阶段是否基本完成，但进入下一阶段前必须先问你确认。", 29, fill=MUTED)
    stages = [
        ("研究设计", "问题、对象、材料边界"),
        ("候选文献", "约 50 篇候选，优先高被引"),
        ("获取全文", "开放来源、Scholar、CNKI、用户文件"),
        ("结构性阅读", "证据句、变量、理论基础"),
        ("引文核验", "PDF 标注、引用是否支撑正文"),
        ("正文写作", "PEEL、章节职责、证据边界"),
        ("图表与 PPT", "论文图、表格、汇报材料"),
        ("审稿修改", "问题账本、回复矩阵、再审再修"),
        ("润色定稿", "目标期刊适配、语言复审"),
        ("投稿复盘", "冻结材料、复现包、经验沉淀"),
    ]
    left = 150
    top = 250
    step_w = 405
    step_h = 145
    for i, (title, desc) in enumerate(stages):
        row = i // 5
        col = i % 5
        x = left + col * 440
        y = top + row * 360
        fill = "#FFFFFF" if i % 2 == 0 else "#F4F7FF"
        draw_round(d, (x, y, x + step_w, y + step_h), fill, "#CFD7E6", 24)
        d.ellipse((x + 28, y + 38, x + 88, y + 98), fill=BLUE if i < 6 else TEAL)
        text(d, (x + 58, y + 51), f"{i + 1}", 24, "#FFFFFF", True, anchor="mm")
        text(d, (x + 112, y + 32), title, 30, bold=True)
        paragraph(d, (x + 112, y + 82), desc, 250, 20)
        if col < 4:
            arrow(d, (x + step_w + 8, y + 74), (x + 432, y + 74), fill="#6B7C93", width=4)
    arrow(d, (left + 4 * 440 + step_w // 2, top + step_h + 40), (left + 4 * 440 + step_w // 2, top + 360 - 24), fill="#6B7C93", width=4)
    arrow(d, (left + 4 * 440 - 20, top + 360 + 74), (left + 3 * 440 + step_w + 8, top + 360 + 74), fill="#6B7C93", width=4)
    arrow(d, (left + 3 * 440 - 20, top + 360 + 74), (left + 2 * 440 + step_w + 8, top + 360 + 74), fill="#6B7C93", width=4)
    arrow(d, (left + 2 * 440 - 20, top + 360 + 74), (left + 1 * 440 + step_w + 8, top + 360 + 74), fill="#6B7C93", width=4)
    arrow(d, (left + 1 * 440 - 20, top + 360 + 74), (left + step_w + 8, top + 360 + 74), fill="#6B7C93", width=4)

    draw_round(d, (280, 1095, 2120, 1280), "#F0FBF8", "#A7DFD4", 30)
    text(d, (330, 1142), "每个阶段都不是自动跳过去", 34, bold=True)
    paragraph(d, (330, 1195), "系统只会提出“建议进入下一阶段”的理由；是否下载、写文件、切换路线、调用多角色或改变项目状态，必须等用户确认。", 1580, 24)
    save(img, path)


def draw_memory_map(path: Path) -> None:
    img = Image.new("RGB", (2400, 1450), PAPER)
    d = ImageDraw.Draw(img)
    text(d, (90, 70), "记忆、自动化与自我演化怎么管", 54, bold=True)
    text(d, (92, 140), "运行态记忆可以让系统越用越顺手，但源规则仍由本地 Git 仓、schema、校验器和提交记录决定。", 29, fill=MUTED)

    draw_round(d, (150, 300, 650, 520), "#FFFFFF", "#CFD7E6", 28)
    text(d, (205, 350), "运行态记忆", 36, bold=True)
    paragraph(d, (205, 405), "agentmemory 记录可召回经验、会话线索和显式保存的偏好。", 360, 23)

    draw_round(d, (150, 650, 650, 870), "#FFFFFF", "#CFD7E6", 28)
    text(d, (205, 700), "代码结构索引", 36, bold=True)
    paragraph(d, (205, 755), "codegraph 帮助理解文件关系和影响范围，只更新忽略缓存。", 360, 23)

    draw_round(d, (900, 420, 1500, 760), "#F4F7FF", "#B9C9EC", 34)
    text(d, (970, 485), "本地源规则层", 44, bold=True)
    rules = ["control_kernel", "routing_table", "conflict_matrix", "schemas", "validators", "tests", "Git commits"]
    for i, item in enumerate(rules):
        x = 970 + (i % 2) * 250
        y = 565 + (i // 2) * 52
        d.rounded_rectangle((x, y, x + 210, y + 34), radius=12, fill="#FFFFFF", outline="#DAE2F4")
        text(d, (x + 14, y + 7), item, 19, fill=BLUE, bold=True)

    draw_round(d, (1750, 300, 2220, 520), "#FFFDF7", "#E7D29C", 28)
    text(d, (1805, 350), "每日检查", 36, bold=True)
    paragraph(d, (1805, 405), "日报扫描路线冲突、七层漂移、记忆状态和演化待办。", 335, 23)

    draw_round(d, (1750, 650, 2220, 870), "#F0FBF8", "#9FDCD1", 28)
    text(d, (1805, 700), "低风险演化", 36, bold=True)
    paragraph(d, (1805, 755), "只处理可回滚、有日志、能通过校验的小改动。", 335, 23)

    arrow(d, (650, 410), (900, 545), fill=BLUE, width=5)
    arrow(d, (650, 760), (900, 650), fill=BLUE, width=5)
    arrow(d, (1500, 545), (1750, 410), fill=ORANGE, width=5)
    arrow(d, (1500, 650), (1750, 760), fill=TEAL, width=5)
    arrow(d, (1985, 870), (1200, 1000), fill="#6B7C93", width=4)
    arrow(d, (1200, 1000), (1200, 760), fill="#6B7C93", width=4)

    draw_round(d, (350, 1080, 2050, 1260), "#FFFFFF", "#D7DDE8", 30)
    text(d, (410, 1130), "硬边界", 34, bold=True)
    paragraph(d, (410, 1182), "记忆召回只提供候选线索；不能静默改 AGENTS、skill、schema、路线表或控制核。所有持久规则必须经过 envctl 校验和 Git 提交。", 1500, 24)
    save(img, path)


def write_svg_asset(path: Path) -> None:
    image2_main = IMAGE2_OVERRIDE_ROOT / "00-local-environment-main-map.png"
    if image2_main.exists():
        svg = """<svg xmlns="http://www.w3.org/2000/svg" width="1600" height="900" viewBox="0 0 1600 900">
  <image href="environment-overview-image2/00-local-environment-main-map.png" x="0" y="0" width="1600" height="900" preserveAspectRatio="xMidYMid meet"/>
</svg>"""
        path.write_text(svg, encoding="utf-8")
        return
    svg = f"""<svg xmlns="http://www.w3.org/2000/svg" width="1600" height="980" viewBox="0 0 1600 980">
  <rect width="1600" height="980" fill="{PAPER}"/>
  <text x="72" y="90" font-family="Microsoft YaHei, Arial" font-size="46" font-weight="700" fill="{INK}">本地研究环境</text>
  <text x="72" y="136" font-family="Microsoft YaHei, Arial" font-size="22" fill="{MUTED}">以工程控制论组织研究、工具、证据和长期演化</text>
  <rect x="130" y="210" width="320" height="120" rx="22" fill="#fff" stroke="#cfd7e6"/>
  <text x="168" y="260" font-family="Microsoft YaHei, Arial" font-size="25" font-weight="700" fill="{INK}">用户任务</text>
  <text x="168" y="296" font-family="Microsoft YaHei, Arial" font-size="17" fill="{MUTED}">研究、文献、写作、分析、PPT</text>
  <path d="M450 270 L560 270" stroke="{BLUE}" stroke-width="5" marker-end="url(#a)"/>
  <rect x="560" y="190" width="360" height="160" rx="26" fill="{BLUE}"/>
  <text x="610" y="250" font-family="Microsoft YaHei, Arial" font-size="28" font-weight="700" fill="#fff">research-autopilot</text>
  <text x="610" y="292" font-family="Microsoft YaHei, Arial" font-size="18" fill="#eaf1ff">先判断路线，再选工具和检查点</text>
  <path d="M920 270 L1040 270" stroke="{BLUE}" stroke-width="5" marker-end="url(#a)"/>
  <rect x="1040" y="210" width="390" height="120" rx="22" fill="#f0fbf8" stroke="#9fdcd1"/>
  <text x="1085" y="260" font-family="Microsoft YaHei, Arial" font-size="25" font-weight="700" fill="{INK}">控制核与规则</text>
  <text x="1085" y="296" font-family="Microsoft YaHei, Arial" font-size="17" fill="{MUTED}">目标、边界、反馈、稳定性和记忆准入</text>
  <g font-family="Microsoft YaHei, Arial" font-size="18">
    {''.join(f'<rect x="{150+(i%4)*340}" y="{430+(i//4)*110}" width="285" height="76" rx="16" fill="#fff" stroke="#d7dde8"/><text x="{175+(i%4)*340}" y="{475+(i//4)*110}" fill="{INK}" font-weight="700">{label}</text>' for i, label in enumerate(["执行动作","工具接口","上下文与证据","研究阶段","运行日志","可靠性检查","环境治理"]))}
  </g>
  <rect x="260" y="760" width="1080" height="94" rx="22" fill="#fffdf7" stroke="#e7d29c"/>
  <text x="310" y="818" font-family="Microsoft YaHei, Arial" font-size="24" font-weight="700" fill="{INK}">工具与知识池：Zotero · Obsidian · agentmemory · codegraph · Git</text>
  <defs><marker id="a" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="{BLUE}"/></marker></defs>
</svg>"""
    path.write_text(svg, encoding="utf-8")


def generate_powerpoint_visuals() -> bool:
    script = SKILLS_ROOT / "scripts" / "generate_environment_overview_visuals.ps1"
    if not script.exists():
        return False
    shell = shutil.which("pwsh") or shutil.which("powershell") or "powershell"
    command = [
        shell,
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(script),
        "-RepoRoot",
        str(REPO_ROOT),
    ]
    result = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace", check=False)
    if result.returncode != 0:
        if result.stdout:
            print(result.stdout)
        if result.stderr:
            print(result.stderr, file=sys.stderr)
        return False
    required = [
        IMAGE_ROOT / "00-local-environment-main-map.png",
        IMAGE_ROOT / "01-research-stage-roadmap.png",
        IMAGE_ROOT / "02-memory-automation-governance.png",
        OUTPUT_ROOT / "local-research-environment-overview.pptx",
    ]
    return all(item.exists() for item in required)


def apply_image2_overrides() -> bool:
    """Use curated image2 visuals when present; they preserve the approved product visual style."""
    names = [
        "00-local-environment-main-map.png",
        "01-research-stage-roadmap.png",
        "02-memory-automation-governance.png",
    ]
    if not all((IMAGE2_OVERRIDE_ROOT / name).exists() for name in names):
        return False
    IMAGE_ROOT.mkdir(parents=True, exist_ok=True)
    presentation_images = PRESENTATION_ROOT / "images"
    presentation_images.mkdir(parents=True, exist_ok=True)
    for name in names:
        source = IMAGE2_OVERRIDE_ROOT / name
        shutil.copy2(source, IMAGE_ROOT / name)
        shutil.copy2(source, presentation_images / name)
    shutil.copy2(IMAGE2_OVERRIDE_ROOT / names[0], ASSET_ROOT / "local-environment-map.png")
    return True


def collect_context() -> dict[str, Any]:
    skills = load_json("skill_catalog.json")["skills"]
    routing = load_json("routing_table.json")["routes"]
    gates = load_json("quality_gates.json")["gates"]
    layers = load_json("environment_layer_contract.json")
    memory = load_json("local_memory_system.json")
    active_skills = {k: v for k, v in skills.items() if v.get("status") == "active"}
    categories: dict[str, list[tuple[str, str]]] = {}
    for name, item in active_skills.items():
        categories.setdefault(item.get("category", "其他"), []).append((name, SKILL_PURPOSES.get(name, f"{item.get('task_type', '专项任务')}。")))
    for values in categories.values():
        values.sort()
    mcp = [item["id"] for item in layers.get("tool_inventory", []) if item.get("kind") == "mcp"]
    return {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "skills": active_skills,
        "skill_categories": categories,
        "routes": routing,
        "gates": gates,
        "layers": layers,
        "memory": memory,
        "mcp": mcp,
        "profiles": sorted(p.stem for p in (SKILLS_ROOT / "profiles").glob("*.toml")),
    }


def md_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = "| " + " | ".join(rows[0]) + " |\n"
    sep = "| " + " | ".join("---" for _ in rows[0]) + " |\n"
    body = "".join("| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |\n" for row in rows[1:])
    return header + sep + body


def build_markdown(ctx: dict[str, Any]) -> str:
    counts = {
        "技能": len(ctx["skills"]),
        "任务路线": len(ctx["routes"]),
        "工具接口": len(ctx["mcp"]),
        "工具组合": len(ctx["profiles"]),
        "检查点": len(ctx["gates"]),
    }
    lines = [
        "# 本地研究环境说明书（更新版）",
        "",
        f"生成时间：{ctx['generated_at']}",
        f"仓库：`{REPO_ROOT}`",
        "",
        "## 一句话",
        "",
        "这套环境是一个受控研究工作台：先判断任务路线，再选择工具和技能；执行后用证据、日志和校验器检查；项目结束后把稳定经验沉淀到规则、技能或长期笔记里。",
        "",
        "## 当前规模",
        "",
        md_table([["项目", "数量"], *[[k, str(v)] for k, v in counts.items()]]),
        "",
        "## 一张图看懂",
        "",
        "![本地研究环境主结构图](images/00-local-environment-main-map.png)",
        "",
        "![研究阶段推进图](images/01-research-stage-roadmap.png)",
        "",
        "![记忆与自动化治理图](images/02-memory-automation-governance.png)",
        "",
        "## 最新结构变化",
        "",
        "- 环境现在按七层管理：执行动作、工具接口、上下文与证据、研究阶段、运行日志、可靠性检查、环境治理。",
        "- `conflict_matrix` 管路线冲突；`route explain` 可以解释为什么命中某条路线，语义模糊时先询问。",
        "- `agentmemory` 是运行态记忆召回层，`codegraph` 是代码结构索引；二者都不能替代本地 Git 仓的源规则。",
        "- 每日自动化会做路线冲突检查、七层结构检查、启动摘要和记忆对账；低风险改动才允许自动落地。",
        "",
        "## 四个核心角色",
        "",
        md_table([
            ["角色", "真实名称", "作用"],
            ["总入口", "research-autopilot", "先判断任务路线，再安排技能、工具、检查点。"],
            ["项目分工", "research-team-orchestrator", "复杂项目进入多角色协作时，明确谁执行、谁审查、交付到哪里。"],
            ["环境治理", "research-stack-manager", "检查技能、插件、工具接口、路径、规则和运行态是否一致。"],
            ["复盘演化", "project-retrospective-evolver", "项目结束后把经验转成可审查的环境演化候选。"],
        ]),
        "",
        "## 日常怎么用",
        "",
        md_table([
            ["你要做什么", "默认进入哪里", "要注意什么"],
            ["筛选参考文献、结构性阅读、引文核验", "evidence-based-literature-workflow", "先建候选表，再获取全文和证据句。"],
            ["下载或整理论文 PDF", "reference-fulltext-acquisition", "英文先开放来源，失败后逐篇 Scholar；中文走授权 CNKI；都归入项目目录。"],
            ["写论文、改稿、润色、目标期刊适配", "manuscript-writing-studio", "不改事实、变量、结果、引用和局限。"],
            ["论文图表、机制图、科研插图", "research-figure-studio / figure-table-studio", "锁定图型、面板、数据和文字后再生成。"],
            ["答辩、汇报、网页 PPT、PPTX", "research-presentation-studio", "默认走 guizang 视觉系统；需要 pptx 时同步输出文件。"],
            ["审稿、内审、返修回复", "academic-paper-review / reviewer-response-pack", "区分致命问题、可修复问题和表层语言问题。"],
            ["环境维护、插件和 MCP 检查", "research-stack-manager", "先扫描和校验，再做最小改动。"],
        ]),
        "",
        "## 技能清单",
        "",
    ]
    for category, items in sorted(ctx["skill_categories"].items()):
        lines.append(f"### {category}")
        lines.append("")
        lines.append(md_table([["名称", "用途"], *[[name, purpose] for name, purpose in items]]))
        lines.append("")
    lines += [
        "## 工具接口清单",
        "",
        md_table([["名称", "用途"], *[[name, MCP_PURPOSES.get(name, "路线限定的工具接口。")] for name in sorted(ctx["mcp"])] ]),
        "",
        "## 常用命令",
        "",
        "```powershell",
        "python -m skills.scripts.envctl route explain \"我要做 revision package\" --summary",
        "python -m skills.scripts.envctl route startup-summary --route-id writing-export --summary",
        "python -m skills.scripts.envctl memory reconcile --summary --probe-agentmemory",
        "python -m skills.scripts.envctl validate conflicts --summary",
        "python -m skills.scripts.envctl validate environment-layers --summary",
        "python -m skills.scripts.envctl validate stack --summary",
        "```",
        "",
        "## 使用边界",
        "",
        "- 正式引用必须可核验；有 DOI 就核验 DOI，没有 DOI 也必须有用户来源、全文 PDF 或公开学术检索证据。",
        "- 新链条、路线切换、高成本工具、多角色协作、下载链和运行态配置变更前，系统应先问用户确认。",
        "- `scholar-nuwa` 是知识蒸馏目录，环境自动化不得写入、同步、清理或覆盖。",
        "- 外部 GitHub 仓库只能通过审查、提案、schema、validator 和测试吸收，不能整包导入替代本地工作流。",
    ]
    return "\n".join(lines) + "\n"


def build_html(ctx: dict[str, Any], markdown_summary: str) -> str:
    category_sections = []
    for category, items in sorted(ctx["skill_categories"].items()):
        rows = "".join(f"<tr><td>{html.escape(name)}</td><td>{html.escape(purpose)}</td></tr>" for name, purpose in items)
        category_sections.append(f"<details><summary>{html.escape(category)}：{len(items)} 个</summary><table><tr><th>名称</th><th>用途</th></tr>{rows}</table></details>")
    mcp_rows = "".join(f"<tr><td>{html.escape(name)}</td><td>{html.escape(MCP_PURPOSES.get(name, '路线限定的工具接口。'))}</td></tr>" for name in sorted(ctx["mcp"]))
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>本地研究环境说明书（更新版）</title>
<style>
:root{{--paper:#fafaf8;--ink:#0a0a0a;--muted:#5f6673;--line:#ddd;--blue:#002fa7;--soft:#f4f7ff;--green:#f0fbf8}}
*{{box-sizing:border-box}}body{{margin:0;background:var(--paper);color:var(--ink);font-family:-apple-system,BlinkMacSystemFont,'Segoe UI','Microsoft YaHei',Arial,sans-serif;line-height:1.72}}
.layout{{display:grid;grid-template-columns:270px 1fr;min-height:100vh}}nav{{position:sticky;top:0;height:100vh;overflow:auto;background:#fff;border-right:1px solid var(--line);padding:28px 22px}}nav a{{display:block;color:#48515f;text-decoration:none;font-size:14px;padding:7px 0}}main{{max-width:1180px;margin:0 auto;padding:54px 54px 90px}}
.hero{{border:1px solid var(--line);background:#fff;border-radius:0;padding:42px 48px;margin-bottom:28px}}.kicker{{font-size:13px;color:var(--blue);font-weight:700;letter-spacing:.12em}}h1{{font-size:56px;line-height:1.05;margin:12px 0 20px;font-weight:300;letter-spacing:0}}h2{{font-size:32px;margin:54px 0 18px;border-top:1px solid var(--line);padding-top:28px}}h3{{font-size:20px}}p{{font-size:16px;color:#2d3440}}.lead{{font-size:20px;color:#1f2937;max-width:880px}}
.metrics{{display:grid;grid-template-columns:repeat(5,1fr);border-top:1px solid var(--line);border-bottom:1px solid var(--line);margin-top:28px}}.metric{{padding:18px;border-right:1px solid var(--line)}}.metric:last-child{{border-right:0}}.metric b{{display:block;font-size:42px;font-weight:300;color:var(--blue)}}.metric span{{font-size:14px;color:var(--muted)}}
.grid{{display:grid;grid-template-columns:repeat(2,1fr);gap:16px}}.card{{background:#fff;border:1px solid var(--line);padding:20px}}figure{{margin:24px 0;padding:18px;background:#fff;border:1px solid var(--line)}}figure img{{max-width:100%;display:block}}figcaption{{font-size:14px;color:var(--muted);margin-top:10px}}table{{width:100%;border-collapse:collapse;background:#fff;margin:14px 0 24px}}th,td{{border:1px solid var(--line);padding:10px 12px;text-align:left;vertical-align:top;font-size:14px}}th{{background:#f0f0ee}}details{{background:#fff;border:1px solid var(--line);padding:0 16px;margin:12px 0}}summary{{cursor:pointer;padding:14px 0;font-weight:700}}code{{background:#f0f0ee;padding:2px 6px}}
@media(max-width:980px){{.layout{{display:block}}nav{{position:relative;height:auto}}main{{padding:28px 18px}}.metrics,.grid{{grid-template-columns:1fr}}h1{{font-size:38px}}table{{display:block;overflow:auto}}}}
</style>
</head>
<body>
<div class="layout"><nav><h3>目录</h3><a href="#top">总览</a><a href="#map">图片</a><a href="#workflow">工作流</a><a href="#skills">技能清单</a><a href="#tools">工具接口</a><a href="#commands">常用命令</a></nav>
<main>
<section class="hero" id="top"><div class="kicker">LOCAL RESEARCH ENVIRONMENT · UPDATED {html.escape(ctx['generated_at'])}</div><h1>本地研究环境说明书</h1><p class="lead">这是一套受控研究工作台：先判断任务路线，再选择工具和技能；执行后用证据、日志和校验器检查；项目结束后把稳定经验沉淀到规则、技能或长期笔记里。</p>
<div class="metrics"><div class="metric"><b>{len(ctx['skills'])}</b><span>技能</span></div><div class="metric"><b>{len(ctx['routes'])}</b><span>任务路线</span></div><div class="metric"><b>{len(ctx['mcp'])}</b><span>工具接口</span></div><div class="metric"><b>{len(ctx['profiles'])}</b><span>工具组合</span></div><div class="metric"><b>{len(ctx['gates'])}</b><span>检查点</span></div></div></section>
<section id="map"><h2>三张图看懂</h2><figure><img src="images/00-local-environment-main-map.png" alt="本地研究环境主结构图"/><figcaption>主结构图：任务先进入总入口，再进入控制核、七层结构、研究工作流、工具与知识池，最后由检查反馈回到总入口。</figcaption></figure><figure><img src="images/01-research-stage-roadmap.png" alt="研究阶段推进图"/><figcaption>阶段推进图：AI 可以建议进入下一阶段，但下载、写入、切换路线和调用多角色前必须确认。</figcaption></figure><figure><img src="images/02-memory-automation-governance.png" alt="记忆与自动化治理图"/><figcaption>记忆治理图：agentmemory 和 codegraph 是辅助层，本地 Git 仓、schema、validator 和提交记录仍是源规则。</figcaption></figure></section>
<section id="workflow"><h2>现在的核心变化</h2><div class="grid"><div class="card"><h3>七层结构</h3><p>执行动作、工具接口、上下文与证据、研究阶段、运行日志、可靠性检查、环境治理。</p></div><div class="card"><h3>路线解释</h3><p><code>route explain</code> 解释一句话为什么进入某条路线；歧义任务先澄清。</p></div><div class="card"><h3>记忆对账</h3><p><code>memory reconcile</code> 对账自动化报告、演化待办和 agentmemory 状态。</p></div><div class="card"><h3>受控演化</h3><p>每日自动化可以发现、登记和检查；持久改动仍要通过校验和 Git。</p></div></div></section>
<section id="skills"><h2>技能清单</h2>{''.join(category_sections)}</section>
<section id="tools"><h2>工具接口</h2><table><tr><th>名称</th><th>用途</th></tr>{mcp_rows}</table></section>
<section id="commands"><h2>常用命令</h2><table><tr><th>目的</th><th>命令</th></tr><tr><td>解释路线</td><td><code>python -m skills.scripts.envctl route explain "我要做 revision package" --summary</code></td></tr><tr><td>生成启动摘要</td><td><code>python -m skills.scripts.envctl route startup-summary --route-id writing-export --summary</code></td></tr><tr><td>记忆对账</td><td><code>python -m skills.scripts.envctl memory reconcile --summary --probe-agentmemory</code></td></tr><tr><td>冲突检查</td><td><code>python -m skills.scripts.envctl validate conflicts --summary</code></td></tr><tr><td>七层检查</td><td><code>python -m skills.scripts.envctl validate environment-layers --summary</code></td></tr></table></section>
</main></div></body></html>"""


def build_docx(ctx: dict[str, Any], md: str, output: Path) -> bool:
    if Document is None or Inches is None:
        return False
    doc = Document()
    doc.add_heading("本地研究环境说明书（更新版）", level=1)
    doc.add_paragraph(f"生成时间：{ctx['generated_at']}")
    doc.add_paragraph("这套环境是一个受控研究工作台：先判断任务路线，再选择工具和技能；执行后用证据、日志和校验器检查；项目结束后把稳定经验沉淀下来。")
    for image in ["00-local-environment-main-map.png", "01-research-stage-roadmap.png", "02-memory-automation-governance.png"]:
        doc.add_picture(str(IMAGE_ROOT / image), width=Inches(6.5))
    doc.add_heading("核心命令", level=2)
    for command in [
        'python -m skills.scripts.envctl route explain "我要做 revision package" --summary',
        "python -m skills.scripts.envctl memory reconcile --summary --probe-agentmemory",
        "python -m skills.scripts.envctl validate conflicts --summary",
        "python -m skills.scripts.envctl validate environment-layers --summary",
    ]:
        doc.add_paragraph(command, style=None)
    doc.save(output)
    return True


EMU = 914400
SLIDE_W = 13.333333
SLIDE_H = 7.5


def emu(v: float) -> int:
    return int(v * EMU)


def ppt_text_box(idx: int, x: float, y: float, w: float, h: float, value: str, size: int, color: str = "0A0A0A", bold: bool = False) -> str:
    paras = "".join(
        f'<a:p><a:r><a:rPr lang="zh-CN" sz="{size * 100}"{" b=\"1\"" if bold else ""}><a:solidFill><a:srgbClr val="{color}"/></a:solidFill><a:latin typeface="Microsoft YaHei"/><a:ea typeface="Microsoft YaHei"/></a:rPr><a:t>{escape(line)}</a:t></a:r></a:p>'
        for line in value.split("\n")
    )
    return f"""<p:sp><p:nvSpPr><p:cNvPr id="{idx}" name="Text {idx}"/><p:cNvSpPr txBox="1"/><p:nvPr/></p:nvSpPr><p:spPr><a:xfrm><a:off x="{emu(x)}" y="{emu(y)}"/><a:ext cx="{emu(w)}" cy="{emu(h)}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:noFill/><a:ln><a:noFill/></a:ln></p:spPr><p:txBody><a:bodyPr wrap="square"/><a:lstStyle/>{paras}</p:txBody></p:sp>"""


def ppt_rect(idx: int, x: float, y: float, w: float, h: float, color: str) -> str:
    return f"""<p:sp><p:nvSpPr><p:cNvPr id="{idx}" name="Rect {idx}"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr><p:spPr><a:xfrm><a:off x="{emu(x)}" y="{emu(y)}"/><a:ext cx="{emu(w)}" cy="{emu(h)}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:solidFill><a:srgbClr val="{color}"/></a:solidFill><a:ln><a:noFill/></a:ln></p:spPr></p:sp>"""


def ppt_image(idx: int, rid: str, x: float, y: float, w: float, h: float) -> str:
    return f"""<p:pic><p:nvPicPr><p:cNvPr id="{idx}" name="Image {idx}"/><p:cNvPicPr/><p:nvPr/></p:nvPicPr><p:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></p:blipFill><p:spPr><a:xfrm><a:off x="{emu(x)}" y="{emu(y)}"/><a:ext cx="{emu(w)}" cy="{emu(h)}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></p:spPr></p:pic>"""


def slide_xml(shapes: list[str]) -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>{''.join(shapes)}</p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sld>"""


def write_pptx(ctx: dict[str, Any], output: Path) -> None:
    slides: list[tuple[list[str], list[Path]]] = []
    images = [
        IMAGE_ROOT / "00-local-environment-main-map.png",
        IMAGE_ROOT / "01-research-stage-roadmap.png",
        IMAGE_ROOT / "02-memory-automation-governance.png",
    ]
    slides.append(([ppt_rect(2, 0, 0, SLIDE_W, SLIDE_H, "002FA7"), ppt_text_box(3, 0.75, 1.25, 11.5, 2.0, "本地研究环境\n一套可治理的研究工作台", 42, "FFFFFF"), ppt_text_box(4, 0.82, 4.55, 10.0, 0.8, f"更新版 · {ctx['generated_at']}", 18, "FFFFFF")], []))
    slides.append(([ppt_text_box(2, 0.65, 0.55, 11.5, 0.7, "一句话", 18, "002FA7", True), ppt_text_box(3, 0.65, 1.2, 11.5, 2.0, "先判断任务路线，再选择工具和技能；执行后检查证据、日志和可复现性；最后把稳定经验沉淀到规则、技能或长期笔记。", 30), ppt_rect(4, 0.65, 4.5, 2.1, 1.0, "F0F0EE"), ppt_text_box(5, 0.85, 4.78, 1.7, 0.4, f"{len(ctx['skills'])} 个技能", 18, "002FA7", True), ppt_rect(6, 3.0, 4.5, 2.1, 1.0, "F0F0EE"), ppt_text_box(7, 3.2, 4.78, 1.7, 0.4, f"{len(ctx['routes'])} 条路线", 18, "002FA7", True), ppt_rect(8, 5.35, 4.5, 2.1, 1.0, "F0F0EE"), ppt_text_box(9, 5.55, 4.78, 1.7, 0.4, f"{len(ctx['mcp'])} 个接口", 18, "002FA7", True)], []))
    slides.append(([ppt_text_box(2, 0.65, 0.38, 10.5, 0.5, "主结构", 18, "002FA7", True), ppt_image(3, "rId2", 0.55, 0.95, 12.2, 6.1)], [images[0]]))
    slides.append(([ppt_text_box(2, 0.65, 0.38, 10.5, 0.5, "研究阶段", 18, "002FA7", True), ppt_image(3, "rId2", 0.55, 0.95, 12.2, 6.1)], [images[1]]))
    slides.append(([ppt_text_box(2, 0.65, 0.38, 10.5, 0.5, "记忆与自动化", 18, "002FA7", True), ppt_image(3, "rId2", 0.55, 0.95, 12.2, 6.1)], [images[2]]))
    slides.append(([ppt_text_box(2, 0.75, 0.75, 6.2, 0.9, "七层结构", 32), ppt_text_box(3, 0.8, 1.9, 5.5, 4.4, "执行动作\n工具接口\n上下文与证据\n研究阶段\n运行日志\n可靠性检查\n环境治理", 24, "002FA7", True), ppt_text_box(4, 6.8, 1.55, 5.4, 3.6, "七层不是为了增加术语，而是让每个组件知道自己负责什么、能写哪里、失败后怎么观察、什么时候必须阻断。", 24)], []))
    slides.append(([ppt_text_box(2, 0.75, 0.75, 7.0, 0.9, "现在新增的增强", 32), ppt_text_box(3, 0.8, 1.95, 11.3, 4.2, "1. conflict_matrix：路线和工具冲突统一登记\n2. route explain：解释为什么命中某条路线\n3. startup-summary：新线程只带紧凑上下文\n4. memory reconcile：对账自动化、待办和运行态记忆", 24)], []))
    slides.append(([ppt_text_box(2, 0.75, 0.75, 7.0, 0.9, "文献、写作、绘图、审稿", 32), ppt_text_box(3, 0.8, 1.95, 11.3, 4.2, "文献获取统一走 reference-fulltext-acquisition。\n结构性阅读和引文核验优先走 evidence-based-literature-workflow。\n科研绘图走 research-figure-studio，PPT 走 research-presentation-studio。\n审稿和返修走 peer review 与 reviewer-response-pack。", 22)], []))
    slides.append(([ppt_text_box(2, 0.75, 0.75, 7.0, 0.9, "使用边界", 32), ppt_text_box(3, 0.8, 1.95, 11.3, 4.2, "正式引用必须可核验。\n新链条和高成本工具先确认。\nagentmemory 与 codegraph 只是辅助层。\nscholar-nuwa 是受保护知识蒸馏目录。\n所有持久改动必须留下校验和 Git 痕迹。", 24)], []))
    slides.append(([ppt_rect(2, 0, 0, SLIDE_W, SLIDE_H, "0A0A0A"), ppt_text_box(3, 0.85, 1.1, 10.5, 1.6, "Open in blue.\nClose in blue.", 44, "FFFFFF"), ppt_text_box(4, 0.9, 5.6, 11.0, 0.8, "研究环境的目标不是多装工具，而是让每一步都可解释、可验证、可回滚。", 22, "FFFFFF")], []))

    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as z:
        slide_count = len(slides)
        overrides = "\n".join(f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>' for i in range(1, slide_count + 1))
        z.writestr("[Content_Types].xml", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Default Extension="png" ContentType="image/png"/><Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>{overrides}<Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/><Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/><Override PartName="/ppt/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/><Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/><Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/></Types>""")
        z.writestr("_rels/.rels", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/><Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/></Relationships>""")
        sld_ids = "".join(f'<p:sldId id="{255+i}" r:id="rId{i}"/>' for i in range(1, slide_count + 1))
        z.writestr("ppt/presentation.xml", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId{slide_count+1}"/></p:sldMasterIdLst><p:sldIdLst>{sld_ids}</p:sldIdLst><p:sldSz cx="{emu(SLIDE_W)}" cy="{emu(SLIDE_H)}" type="wide"/><p:notesSz cx="6858000" cy="9144000"/></p:presentation>""")
        rels = "".join(f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>' for i in range(1, slide_count + 1))
        rels += f'<Relationship Id="rId{slide_count+1}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="slideMasters/slideMaster1.xml"/>'
        z.writestr("ppt/_rels/presentation.xml.rels", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{rels}</Relationships>""")
        z.writestr("ppt/slideMasters/slideMaster1.xml", slide_xml([]))
        z.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="../theme/theme1.xml"/></Relationships>""")
        z.writestr("ppt/slideLayouts/slideLayout1.xml", slide_xml([]))
        z.writestr("ppt/slideLayouts/_rels/slideLayout1.xml.rels", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="../slideMasters/slideMaster1.xml"/></Relationships>""")
        z.writestr("ppt/theme/theme1.xml", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Local Research Environment"><a:themeElements><a:clrScheme name="Office"><a:dk1><a:srgbClr val="0A0A0A"/></a:dk1><a:lt1><a:srgbClr val="FAFAF8"/></a:lt1><a:dk2><a:srgbClr val="002FA7"/></a:dk2><a:lt2><a:srgbClr val="F0F0EE"/></a:lt2><a:accent1><a:srgbClr val="002FA7"/></a:accent1><a:accent2><a:srgbClr val="0F9F8F"/></a:accent2><a:accent3><a:srgbClr val="F59E0B"/></a:accent3><a:accent4><a:srgbClr val="6E56CF"/></a:accent4><a:accent5><a:srgbClr val="17A673"/></a:accent5><a:accent6><a:srgbClr val="64748B"/></a:accent6><a:hlink><a:srgbClr val="002FA7"/></a:hlink><a:folHlink><a:srgbClr val="002FA7"/></a:folHlink></a:clrScheme><a:fontScheme name="Office"><a:majorFont><a:latin typeface="Microsoft YaHei"/><a:ea typeface="Microsoft YaHei"/></a:majorFont><a:minorFont><a:latin typeface="Microsoft YaHei"/><a:ea typeface="Microsoft YaHei"/></a:minorFont></a:fontScheme><a:fmtScheme name="Office"><a:fillStyleLst/><a:lnStyleLst/><a:effectStyleLst/><a:bgFillStyleLst/></a:fmtScheme></a:themeElements></a:theme>""")
        z.writestr("docProps/core.xml", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>本地研究环境介绍</dc:title><dc:creator>Codex</dc:creator><cp:lastModifiedBy>Codex</cp:lastModifiedBy><dcterms:created xsi:type="dcterms:W3CDTF">{datetime.now(UTC).isoformat().replace('+00:00', 'Z')}</dcterms:created></cp:coreProperties>""")
        z.writestr("docProps/app.xml", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><Application>Codex</Application><PresentationFormat>On-screen Show (16:9)</PresentationFormat></Properties>""")
        media_index = 1
        for i, (shapes, slide_images) in enumerate(slides, start=1):
            z.writestr(f"ppt/slides/slide{i}.xml", slide_xml(shapes))
            rel_items = ['<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/>']
            for j, image_path in enumerate(slide_images, start=2):
                target = f"../media/image{media_index}.png"
                rel_items.append(f'<Relationship Id="rId{j}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="{target}"/>')
                z.write(image_path, f"ppt/media/image{media_index}.png")
                media_index += 1
            z.writestr(f"ppt/slides/_rels/slide{i}.xml.rels", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{''.join(rel_items)}</Relationships>""")


def build_html_deck(ctx: dict[str, Any], output: Path) -> None:
    template = (GUIZANG_ROOT / "assets" / "template-swiss.html").read_text(encoding="utf-8")
    image_dir = output.parent / "images"
    image_dir.mkdir(parents=True, exist_ok=True)
    for image in ["00-local-environment-main-map.png", "01-research-stage-roadmap.png", "02-memory-automation-governance.png"]:
        shutil.copy2(IMAGE_ROOT / image, image_dir / image)
    slides = f"""
<section class="slide accent" data-layout="SWISS-COVER-ASCII" data-animate="hero"><div class="canvas-card"><div class="chrome-min">LOCAL RESEARCH ENVIRONMENT · UPDATED</div><h1 style="align-self:center;font-family:var(--sans),var(--sans-zh);font-weight:200;font-size:min(8.4vw,15vh);line-height:.96;color:#fff">本地研究环境<br/>可治理的研究工作台</h1><div class="foot">Engineering Cybernetics · Seven Layers · Controlled Evolution</div></div></section>
<section class="slide light" data-layout="S09"><div class="canvas-card"><div class="chrome-min">SECTION 01 · IDEA</div><div class="t-meta">ONE SENTENCE</div><h2 class="h-statement">先判断路线<br/>再调用工具<br/>最后留下证据。</h2><div class="foot">不是堆工具，而是让每一步可解释、可验证、可回滚。</div></div></section>
<section class="slide light" data-layout="S22"><div class="canvas-card"><div class="chrome-min">SECTION 02 · STRUCTURE</div><div class="frame-img r-21x9"><img src="images/00-local-environment-main-map.png" data-image-slot="s22-hero-21x9" style="object-fit:contain;object-position:center center"/></div><div class="kpi-row-3"><div><b>01</b><span>总入口</span></div><div><b>07</b><span>七层结构</span></div><div><b>闭环</b><span>反馈纠偏</span></div></div></div></section>
<section class="slide light" data-layout="S05"><div class="canvas-card"><div class="chrome-min">SECTION 03 · LAYERS</div><div class="t-meta">SEVEN LAYERS</div><h2 class="h-xl-zh">把复杂环境拆成七层</h2><div class="stack-row"><div class="stack-block b-accent"><div class="nb">01</div><h3>执行 / 工具</h3><p>命令、文件、浏览器、Zotero、学术检索。</p></div><div class="stack-block"><div class="nb">02</div><h3>上下文 / 阶段</h3><p>项目事实、证据账本、研究阶段和下一步确认。</p></div><div class="stack-block b-ink"><div class="nb">03</div><h3>日志 / 校验 / 治理</h3><p>报告、schema、质量门、回滚和演化边界。</p></div></div></div></section>
<section class="slide light" data-layout="S22"><div class="canvas-card"><div class="chrome-min">SECTION 04 · WORKFLOW</div><div class="frame-img r-21x9"><img src="images/01-research-stage-roadmap.png" data-image-slot="s22-hero-21x9" style="object-fit:contain;object-position:center center"/></div><div class="foot">阶段推进由 AI 主动建议，但切换路线、下载、写入和多角色协作前必须确认。</div></div></section>
<section class="slide light" data-layout="S08"><div class="canvas-card"><div class="chrome-min">SECTION 05 · SPLIT</div><div class="duo-compare"><div class="col accent"><div class="col-tag"><span class="num">A</span>源规则层</div><h3 class="col-ttl">本地 Git 仓</h3><ul class="col-list"><li>control kernel</li><li>routing table</li><li>schema / validator</li><li>Git commit</li></ul></div><div class="vrule"></div><div class="col"><div class="col-tag"><span class="num">B</span>运行辅助层</div><h3 class="col-ttl">agentmemory / codegraph</h3><ul class="col-list"><li>召回经验</li><li>索引代码关系</li><li>形成候选</li><li>不能直接改规则</li></ul></div></div></div></section>
<section class="slide light" data-layout="S22"><div class="canvas-card"><div class="chrome-min">SECTION 06 · MEMORY</div><div class="frame-img r-21x9"><img src="images/02-memory-automation-governance.png" data-image-slot="s22-hero-21x9" style="object-fit:contain;object-position:center center"/></div><div class="foot">运行态记忆让环境更顺手；持久规则仍由本地合同和校验器决定。</div></div></section>
<section class="slide light" data-layout="S19"><div class="canvas-card"><div class="chrome-min">SECTION 07 · DAILY USE</div><div class="t-meta">HOW TO USE</div><h2 class="h-xl-zh">日常只记四个入口</h2><div class="four-cards"><div><b>研究总入口</b><p>research-autopilot</p></div><div><b>文献证据</b><p>evidence-based-literature-workflow</p></div><div><b>写作图表</b><p>manuscript / figure / presentation</p></div><div><b>环境治理</b><p>research-stack-manager</p></div></div></div></section>
<section class="slide light" data-layout="S10"><div class="canvas-card"><div class="split-half"><div class="half b-accent"><div class="chrome-min">CLOSING</div><h2 style="font-weight:200;font-size:min(7vw,12vh);line-height:.95;color:#fff">让系统读自己的输出<br/>改自己的 skill。</h2></div><div class="half"><div class="chrome-min">NEXT</div><ol style="font-size:2vw;line-height:1.8"><li>路线先解释</li><li>证据先核验</li><li>记忆先对账</li><li>提交前校验</li></ol></div></div></div></section>
"""
    html_text = template.replace("[必填] 替换为 PPT 标题 · Deck Title", "本地研究环境介绍 · Local Research Environment")
    marker = "<!-- SLIDES_HERE"
    nav = '<div id="nav"></div>'
    marker_start = html_text.index(marker)
    nav_start = html_text.index(nav, marker_start)
    html_text = html_text[:marker_start] + slides + "\n\n</div>\n\n" + html_text[nav_start:]
    output.write_text(html_text, encoding="utf-8")


def main() -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    IMAGE_ROOT.mkdir(parents=True, exist_ok=True)
    ASSET_ROOT.mkdir(parents=True, exist_ok=True)
    ctx = collect_context()

    powerpoint_visuals_ok = generate_powerpoint_visuals()
    if not powerpoint_visuals_ok:
        draw_main_map(IMAGE_ROOT / "00-local-environment-main-map.png")
        draw_stage_map(IMAGE_ROOT / "01-research-stage-roadmap.png")
        draw_memory_map(IMAGE_ROOT / "02-memory-automation-governance.png")
        shutil.copy2(IMAGE_ROOT / "00-local-environment-main-map.png", ASSET_ROOT / "local-environment-map.png")
    image2_overrides_ok = apply_image2_overrides()
    write_svg_asset(ASSET_ROOT / "local-environment-map.svg")

    md = build_markdown(ctx)
    (OUTPUT_ROOT / "local-research-environment-guide.md").write_text(md, encoding="utf-8")
    (OUTPUT_ROOT / "local-research-environment-guide.html").write_text(build_html(ctx, md), encoding="utf-8")
    docx_ok = build_docx(ctx, md, OUTPUT_ROOT / "local-research-environment-guide.docx")

    deck_html = PRESENTATION_ROOT / "index.html"
    deck_html.parent.mkdir(parents=True, exist_ok=True)
    build_html_deck(ctx, deck_html)
    if image2_overrides_ok or not powerpoint_visuals_ok:
        write_pptx(ctx, OUTPUT_ROOT / "local-research-environment-overview.pptx")

    manifest = {
        "generated_at": ctx["generated_at"],
        "version": "environment-overview-v4-seven-layer-memory",
        "outputs": {
            "markdown": str(OUTPUT_ROOT / "local-research-environment-guide.md"),
            "html": str(OUTPUT_ROOT / "local-research-environment-guide.html"),
            "docx": str(OUTPUT_ROOT / "local-research-environment-guide.docx") if docx_ok else None,
            "pptx": str(OUTPUT_ROOT / "local-research-environment-overview.pptx"),
            "web_deck": str(deck_html),
        },
        "images": [
            str(IMAGE_ROOT / "00-local-environment-main-map.png"),
            str(IMAGE_ROOT / "01-research-stage-roadmap.png"),
            str(IMAGE_ROOT / "02-memory-automation-governance.png"),
            str(ASSET_ROOT / "local-environment-map.png"),
            str(ASSET_ROOT / "local-environment-map.svg"),
        ],
        "counts": {
            "skills": len(ctx["skills"]),
            "routes": len(ctx["routes"]),
            "mcp": len(ctx["mcp"]),
            "profiles": len(ctx["profiles"]),
            "quality_gates": len(ctx["gates"]),
        },
        "visual_system": "Curated image2 product diagrams with deterministic PowerPoint/PPTX fallback",
        "powerpoint_visuals": powerpoint_visuals_ok,
        "image2_overrides": image2_overrides_ok,
        "source_contracts": [
            "skills/catalog/environment_layer_contract.json",
            "skills/catalog/local_memory_system.json",
            "skills/catalog/conflict_matrix.json",
            "skills/catalog/routing_table.json",
            "skills/catalog/skill_catalog.json",
        ],
    }
    (OUTPUT_ROOT / "local-research-environment-guide.manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (IMAGE_ROOT / "diagram-manifest.json").write_text(json.dumps({"generated_at": ctx["generated_at"], "images": manifest["images"]}, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
